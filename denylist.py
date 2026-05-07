from __future__ import annotations

from pathlib import Path
from typing import List
from docx import Document


def _safe_str(x) -> str:
    if x is None:
        return ""
    s = str(x).strip()
    return s


def load_denylist(path: str) -> List[str]:
    """
    Leest een DOCX met klanten die niet benaderd willen worden.
    We nemen alle niet-lege regels/paragraphs mee.
    """
    p = Path(path)
    if not p.exists():
        return []

    doc = Document(str(p))
    items: List[str] = []
    for para in doc.paragraphs:
        t = _safe_str(para.text)
        if t:
            items.append(t)

    # ook tabelcellen meenemen (soms staan namen in een tabel)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = _safe_str(cell.text)
                if t:
                    items.append(t)

    # uniek + netjes
    uniq = []
    seen = set()
    for it in items:
        key = it.lower()
        if key not in seen:
            seen.add(key)
            uniq.append(it)
    return uniq


def load_denylist_docx(path: str) -> List[str]:
    """
    Backwards compatible alias.
    """
    return load_denylist(path)
